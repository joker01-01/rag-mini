# -*- coding: utf-8 -*-
"""生成演示用测试文档（PDF + DOCX + XLSX），供最小 RAG 系统跑通全流程。

真实使用时把 docs/ 下的文件换成你自己的资料即可，格式不限。
"""

from pathlib import Path

import openpyxl
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

OUT = Path(__file__).parent / "docs"

PRODUCT_MD = [
    ("# 轻雀知识库 产品手册", "h1"),
    ("轻雀知识库是一个面向个人与团队的智能文档管理工具，支持把 PDF、Word、PPT、Excel 等文件统一导入，并通过 AI 问答检索资料。", "body"),
    ("## 一、产品介绍", "h2"),
    ("轻雀知识库主打“检索增强生成”（RAG）技术：先把资料统一转换为 Markdown 并建立索引，回答问题时先检索相关资料，再生成带引用来源的答案。", "body"),
    ("产品适用于产品手册管理、论文阅读、会议纪要归档、客户资料检索等场景。", "body"),
    ("## 二、核心功能", "h2"),
    ("### 1. 多格式导入", "h3"),
    ("支持 PDF、DOCX、PPTX、XLSX、图片、网页链接，单文件最大 50MB。", "body"),
    ("### 2. 智能问答", "h3"),
    ("基于 RAG 技术回答资料相关问题，回答末尾自动标注引用来源，方便溯源核对。", "body"),
    ("### 3. 权限管理", "h3"),
    ("支持四种成员角色：所有者、管理员、编辑者、只读。所有者拥有全部权限，可删除知识库。", "body"),
    ("## 三、系统要求", "h2"),
    ("客户端支持 Windows 10/11、macOS 12 及以上；推荐内存 8GB 以上。", "body"),
    ("离线模式需要在首次使用时下载本地模型（约 2GB），下载完成后即可在没有网络的情况下问答。", "body"),
    ("## 四、定价", "h2"),
    ("免费版：100 个文档，每天 20 次问答。", "body"),
    ("专业版：99 元/月，无限文档，每天 2000 次问答。", "body"),
    ("团队版：299 元/月，包含 5 个成员，支持私有化部署。", "body"),
    ("## 五、退款政策", "h2"),
    ("订阅后 7 天内可申请全额退款；超过 7 天按剩余天数比例退款。", "body"),
    ("退款原路返回，到账时间一般为 3~5 个工作日。", "body"),
]

GUIDE_MD = [
    ("# 轻雀知识库 使用指南", "h1"),
    ("## 一、快速上手", "h2"),
    ("### 1. 安装与登录", "h3"),
    ("在官网下载客户端，注册账号后登录。首次登录建议先创建一个知识库。", "body"),
    ("### 2. 导入第一批文档", "h3"),
    ("点击“导入文档”，把需要检索的 PDF、Word、PPT、Excel 拖入即可。系统会自动完成转换与索引。", "body"),
    ("## 二、常用操作", "h2"),
    ("### 1. 创建知识库", "h3"),
    ("一个知识库对应一个主题，比如“产品资料”、“论文阅读”。“创建知识库”按钮位于首页左上角。", "body"),
    ("### 2. 上传与更新文档", "h3"),
    ("文档更新后，重新拖入同名文件即可覆盖旧版本，索引会自动重建。", "body"),
    ("### 3. 发起问答", "h3"),
    ("在问答框输入问题，系统会在当前知识库内检索相关片段，生成带引用的回答。", "body"),
    ("## 三、数据导入建议", "h2"),
    ("### 1. 文件命名规范", "h3"),
    ("建议使用“主题+日期”的命名方式，例如“产品手册-20260101.pdf”，便于检索时识别来源。", "body"),
    ("### 2. 推荐格式", "h3"),
    ("文字版 PDF 优于扫描版 PDF；扫描件请先做 OCR 再导入，否则内容无法被检索。", "body"),
    ("### 3. 更新频率", "h3"),
    ("建议资料有变化时及时更新，保持知识库与真实资料一致，回答准确率更高。", "body"),
    ("## 四、常见限制", "h2"),
    ("单文件最大 50MB；单个知识库最多 5000 个文档。", "body"),
    ("每次问答最多引用 5 个资料片段，超出部分不会被用于生成回答。", "body"),
]

FAQ_ROWS = [
    ("问题", "答案", "分类"),
    ("导入失败怎么办", "检查文件格式与大小限制（单文件不超过 50MB）；扫描版 PDF 需先做 OCR 再导入。", "导入"),
    ("忘记密码", "在登录页点击“忘记密码”，通过注册邮箱重置。", "账号"),
    ("如何取消订阅", "设置-订阅管理-取消订阅；取消后当前计费周期结束前仍可正常使用。", "计费"),
    ("是否支持离线", "支持。首次使用时下载本地模型（约 2GB），之后可离线问答。", "功能"),
    ("数据是否安全", "文档默认加密存储；团队版以上支持私有化部署，数据不出企业内网。", "安全"),
]


def register_pdf_font() -> str:
    """注册一个支持中文的 PDF 字体，优先 TTF，其次 CID 字体。"""
    for name, path in [("SimHei", r"C:\Windows\Fonts\simhei.ttf"), ("MSYH", r"C:\Windows\Fonts\msyh.ttc")]:
        if Path(path).exists():
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                return name
            except Exception:
                continue
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light"


def make_pdf():
    font = register_pdf_font()
    h1 = ParagraphStyle("h1", fontName=font, fontSize=18, leading=24, spaceAfter=10)
    h2 = ParagraphStyle("h2", fontName=font, fontSize=14, leading=20, spaceBefore=10, spaceAfter=6)
    h3 = ParagraphStyle("h3", fontName=font, fontSize=12, leading=18, spaceBefore=6, spaceAfter=4)
    body = ParagraphStyle("body", fontName=font, fontSize=11, leading=17, spaceAfter=6)
    style_map = {"h1": h1, "h2": h2, "h3": h3, "body": body}

    doc = SimpleDocTemplate(
        str(OUT / "轻雀知识库产品手册.pdf"),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title="轻雀知识库 产品手册",
    )
    story = []
    for text, kind in PRODUCT_MD:
        story.append(Paragraph(text, style_map[kind]))
        story.append(Spacer(1, 2))
    doc.build(story)


def make_docx():
    doc = Document()
    doc.add_heading("轻雀知识库 使用指南", 0)
    current_h2 = None
    for text, kind in GUIDE_MD:
        if kind == "h1":
            continue
        if kind == "h2":
            current_h2 = text.lstrip("#").strip()
            doc.add_heading(current_h2, level=1)
        elif kind == "h3":
            doc.add_heading(text.lstrip("#").strip(), level=2)
        else:
            doc.add_paragraph(text)
    doc.save(str(OUT / "轻雀知识库使用指南.docx"))


def make_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "常见问题"
    for row in FAQ_ROWS:
        ws.append(row)
    widths = [18, 70, 8]
    for idx, w in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(idx)].width = w
    wb.save(str(OUT / "轻雀知识库常见问题.xlsx"))


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    make_pdf()
    make_docx()
    make_xlsx()
    for p in sorted(OUT.iterdir()):
        print(f"生成: {p.name} ({p.stat().st_size} bytes)")
